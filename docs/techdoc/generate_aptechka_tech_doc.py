from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs" / "techdoc"
ASSETS_DIR = DOCS_DIR / "assets"
OUTPUT_DOCX = DOCS_DIR / "aptechka-technical-documentation-defense.docx"
COVER_ICON_PATH = ROOT / "mobile" / "assets" / "app" / "icon.png"
ARCHITECTURE_PATH = ASSETS_DIR / "aptechka-architecture.png"
BUILD_DATE = "12 мая 2026 г."

ACCENT = RGBColor(0x31, 0xBA, 0xCB)
DARK = RGBColor(0x10, 0x27, 0x3C)
MUTED = RGBColor(0x5B, 0x67, 0x75)
LIGHT_BG = "EAF9FB"
LINE = "D9E2EC"
OK_FILL = "EAF7EE"
WARN_FILL = "FFF4E5"
PARTIAL_FILL = "FFF8DD"
INFO_FILL = "EDF5FF"
HEADER_FILL = "10273C"
ACCENT_FILL = "31BACB"


@dataclass(frozen=True)
class TableRow:
    columns: tuple[str, ...]


def pt(value: float) -> Pt:
    return Pt(value)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[index:index + 2], 16) for index in (0, 2, 4))


def ensure_dirs() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def load_font(name: str, size: int) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        probe = f"{current} {word}"
        if draw.textbbox((0, 0), probe, font=font)[2] <= width:
            current = probe
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: tuple[int, int, int],
    width: int,
    line_gap: int = 6,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, font, width)
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line, font=font)
        y += (bbox[3] - bbox[1]) + line_gap
    return y


def rounded_box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    *,
    fill: str,
    outline: str,
    title_fill: tuple[int, int, int] = (16, 39, 60),
    body_fill: tuple[int, int, int] = (64, 76, 91),
) -> None:
    draw.rounded_rectangle(rect, radius=26, fill=fill, outline=outline, width=3)
    title_font = load_font("arialbd.ttf", 30)
    body_font = load_font("arial.ttf", 22)
    left, top, right, _ = rect
    draw.text((left + 28, top + 22), title, font=title_font, fill=title_fill)
    y = top + 72
    for line in lines:
        y = draw_multiline(draw, (left + 28, y), line, body_font, body_fill, right - left - 56, line_gap=4)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    fill: str = "#23748B",
    width: int = 7,
) -> None:
    draw.line([start, end], fill=fill, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    if dx == 0 and dy == 0:
        return
    if abs(dx) >= abs(dy):
        direction = 1 if dx > 0 else -1
        points = [
            end,
            (end[0] - 22 * direction, end[1] - 11),
            (end[0] - 22 * direction, end[1] + 11),
        ]
    else:
        direction = 1 if dy > 0 else -1
        points = [
            end,
            (end[0] - 11, end[1] - 22 * direction),
            (end[0] + 11, end[1] - 22 * direction),
        ]
    draw.polygon(points, fill=fill)


def add_caption(draw: ImageDraw.ImageDraw, x: int, y: int, text: str) -> None:
    font = load_font("arial.ttf", 19)
    draw.text((x, y), text, font=font, fill=(35, 116, 139))


def create_architecture_diagram() -> None:
    image = Image.new("RGB", (1600, 950), color=(248, 252, 253))
    draw = ImageDraw.Draw(image)

    title_font = load_font("arialbd.ttf", 42)
    subtitle_font = load_font("arial.ttf", 22)
    draw.text((70, 44), "Архитектура проекта Aptechka", font=title_font, fill=hex_to_rgb("10273C"))
    draw.text(
        (70, 102),
        "Система состоит из веб-клиента, мобильного клиента, Supabase и отдельного push-worker.",
        font=subtitle_font,
        fill=hex_to_rgb("5B6775"),
    )

    rounded_box(
        draw,
        (70, 180, 470, 390),
        "Веб-клиент",
        [
            "React 19 + React Router",
            "Лендинг, авторизация, инвентарь, календарь, приглашения, профиль",
            "Service Worker, web push, локальные браузерные напоминания",
        ],
        fill="#EDF9FC",
        outline="#31BACB",
    )

    rounded_box(
        draw,
        (70, 450, 470, 700),
        "Мобильный клиент",
        [
            "Expo 54 + React Native 0.81",
            "Экраны входа, аптечек, добавления лекарств, расписания и push",
            "Expo Notifications, единый Supabase backend",
        ],
        fill="#F7FBFF",
        outline="#7BB6F8",
    )

    rounded_box(
        draw,
        (570, 160, 1040, 735),
        "Supabase",
        [
            "Auth: регистрация и сессии пользователей",
            "Postgres: аптечки, лекарства, расписания, журналы, приглашения, уведомления",
            "Storage: avatars, medicine-images",
            "RLS: таблица mobile_push_tokens явно настроена в SQL из репозитория",
        ],
        fill="#FFFFFF",
        outline="#10273C",
    )

    rounded_box(
        draw,
        (1135, 155, 1530, 410),
        "Push-worker",
        [
            "Node 18-alpine",
            "Опрашивает medicine_schedules по времени и таймзоне",
            "Отправляет Web Push и Expo Push, пишет push_notification_logs",
        ],
        fill="#FFF8EC",
        outline="#F0B14A",
    )

    rounded_box(
        draw,
        (1135, 490, 1530, 725),
        "Каналы доставки",
        [
            "Web Push + VAPID для браузеров",
            "Expo Push API для мобильных устройств",
            "Deep link /calendar и screen=calendar",
        ],
        fill="#F4FDF7",
        outline="#4EA971",
    )

    rounded_box(
        draw,
        (570, 790, 1530, 905),
        "Развёртывание",
        [
            "Docker Compose поднимает web и push-worker, web обслуживается через nginx и сертификаты Let's Encrypt.",
        ],
        fill="#F8FAFC",
        outline="#CBD5E1",
    )

    arrow(draw, (470, 285), (570, 285))
    arrow(draw, (470, 575), (570, 575))
    arrow(draw, (1040, 285), (1135, 285))
    arrow(draw, (1330, 410), (1330, 490))
    arrow(draw, (1000, 735), (1000, 790))
    arrow(draw, (260, 390), (260, 450), fill="#31BACB")

    add_caption(draw, 165, 400, "обновление локальных браузерных уведомлений")
    add_caption(draw, 1180, 440, "формирование и отправка напоминаний")

    image.save(ARCHITECTURE_PATH)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2EC", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in (
        ("Title", 24, DARK),
        ("Subtitle", 13, MUTED),
        ("Heading 1", 16, DARK),
        ("Heading 2", 12.5, DARK),
        ("Heading 3", 11.5, DARK),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = color
        style.font.size = pt(size)

    doc.core_properties.title = "Техническая документация проекта Aptechka"
    doc.core_properties.subject = "Документ для защиты проекта"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.comments = "Сформировано по фактическому состоянию репозитория на 12.05.2026"


def configure_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Aptechka · Техническая документация")
    run.font.name = "Calibri"
    run.font.size = pt(9)
    run.font.color.rgb = MUTED


def add_spacer(doc: Document, size_pt: float = 4) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = pt(size_pt)
    paragraph.paragraph_format.space_before = pt(0)


def add_section_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK


def add_subsection_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT


def add_body(doc: Document, text: str, *, first_line_cm: float | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_cm)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = pt(11)
    run.font.color.rgb = DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(f"- {item}")
        run.font.name = "Calibri"
        run.font.size = pt(11)
        run.font.color.rgb = DARK


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(f"{index}. {item}")
        run.font.name = "Calibri"
        run.font.size = pt(11)
        run.font.color.rgb = DARK


def add_callout(doc: Document, title: str, text: str, fill: str = INFO_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.8)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color="C7D2E0")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = pt(11)
    title_run.font.color.rgb = DARK
    body_paragraph = cell.add_paragraph(text)
    body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_run = body_paragraph.runs[0]
    body_run.font.name = "Calibri"
    body_run.font.size = pt(10.5)
    body_run.font.color.rgb = DARK
    add_spacer(doc, 2)


def configure_table(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False


def style_row_cells(row, widths_cm: list[float]) -> None:
    for index, width in enumerate(widths_cm):
        if index < len(row.cells):
            row.cells[index].width = Cm(width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(row.cells[index], color=LINE)


def style_header_row(row) -> None:
    for cell in row.cells:
        shade_cell(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not paragraph.runs:
            paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_repeat_table_header(row)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    color: RGBColor = DARK,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_two_col_table(doc: Document, rows: list[TableRow], widths_cm: list[float], header: tuple[str, str] | None = None) -> None:
    table = doc.add_table(rows=0, cols=2)
    configure_table(table, widths_cm)
    if header:
        header_row = table.add_row()
        style_row_cells(header_row, widths_cm)
        header_row.cells[0].text = header[0]
        header_row.cells[1].text = header[1]
        style_header_row(header_row)
    for row in rows:
        table_row = table.add_row()
        style_row_cells(table_row, widths_cm)
        set_cell_text(table_row.cells[0], row.columns[0], bold=True, size=10.2)
        set_cell_text(table_row.cells[1], row.columns[1], size=10.2)
    add_spacer(doc, 2)


def add_multi_col_table(doc: Document, headers: list[str], rows: list[TableRow], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    configure_table(table, widths_cm)
    style_row_cells(table.rows[0], widths_cm)
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    style_header_row(table.rows[0])

    for row in rows:
        table_row = table.add_row()
        style_row_cells(table_row, widths_cm)
        for index, value in enumerate(row.columns):
            align = WD_ALIGN_PARAGRAPH.LEFT
            if index in (1, 2) and len(headers) == 4:
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table_row.cells[index], value, size=9.8, align=align)
    add_spacer(doc, 2)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_footer(section)

    top_paragraph = doc.add_paragraph()
    top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_paragraph.add_run().add_picture(str(COVER_ICON_PATH), width=Inches(1.1))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = pt(8)
    title.paragraph_format.space_after = pt(6)
    title_run = title.add_run("Техническая документация проекта\n«Цифровая аптечка»")
    title_run.bold = True
    title_run.font.color.rgb = DARK
    title_run.font.size = pt(24)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = pt(20)
    subtitle_run = subtitle.add_run(
        "Версия для защиты. Система учета лекарств, семейного доступа и автоматических напоминаний."
    )
    subtitle_run.font.color.rgb = MUTED
    subtitle_run.font.size = pt(12.5)

    passport_rows = [
        TableRow(("Проект", "Aptechka / «Цифровая аптечка»")),
        TableRow(("Состояние", f"Фактическое состояние репозитория на {BUILD_DATE}")),
        TableRow(("Платформы", "Web SPA (React) и mobile client (Expo / React Native)")),
        TableRow(("Backend", "Supabase Auth, Postgres, Storage, push-интеграции")),
        TableRow(("Развёртывание", "Docker Compose: web + push-worker")),
    ]
    add_two_col_table(doc, passport_rows, [4.6, 12.2], header=("Параметр", "Описание"))

    add_callout(
        doc,
        "Ключевая мысль проекта",
        "Текущая версия продукта решает практическую задачу ведения домашних и совместных аптечек: хранение состава, контроль расписания приема лекарств, фиксация факта приема и доставку напоминаний в веб и mobile.",
        fill=LIGHT_BG,
    )

    add_body(
        doc,
        "Документ подготовлен по исходному коду проекта, конфигурационным файлам, SQL-скрипту mobile push setup и результатам локальной проверки сборок. Это важно для защиты: ниже описана не абстрактная концепция, а именно то состояние, которое подтверждается текущим репозиторием.",
    )

    doc.add_page_break()


def add_overview(doc: Document) -> None:
    add_section_title(doc, "1. Назначение системы и границы реализации")
    add_body(
        doc,
        "Aptechka представляет собой клиент-серверную систему для персонального и семейного учета лекарств. Пользователь может зарегистрироваться, создать одну или несколько аптечек, добавить состав лекарств, настроить расписание приема и получать напоминания на веб- и мобильных устройствах.",
    )
    add_bullets(
        doc,
        [
            "Центральная сущность проекта — аптечка, которая может быть личной или совместной.",
            "Лекарства относятся к выбранной аптечке, могут иметь дозировку, количество, категорию, срок годности и изображение.",
            "Напоминания формируются по расписанию приема и могут доставляться в браузер и на мобильные устройства.",
            "Система поддерживает приглашения пользователей в совместные аптечки и базовый личный профиль.",
        ],
    )
    add_callout(
        doc,
        "Важно для защиты",
        "Маркетинговые блоки лендинга упоминают запись к врачу и другие расширенные возможности, однако в текущем коде основная предметная область сосредоточена на аптечках, лекарствах, расписаниях и уведомлениях. На защите лучше акцентировать именно это ядро.",
        fill=PARTIAL_FILL,
    )

    add_subsection_title(doc, "Технологический стек")
    add_multi_col_table(
        doc,
        ["Слой", "Инструменты", "Назначение"],
        [
            TableRow(("Веб-клиент", "React 19.2, react-router-dom 7.13, react-scripts 5.0.1", "SPA с лендингом, инвентарем, календарем и профилем")),
            TableRow(("Мобильный клиент", "Expo 54, React Native 0.81, expo-notifications", "Мобильные сценарии учета и push-напоминаний")),
            TableRow(("Backend", "@supabase/supabase-js 2.86.2, Supabase Auth/Postgres/Storage", "Аутентификация, данные, файлы, клиентский API")),
            TableRow(("Push-доставка", "web-push 3.6.7, Expo Push API", "Отправка удаленных напоминаний")),
            TableRow(("Инфраструктура", "Docker Compose 3.8, nginx, Node 18-alpine", "Сборка и публикация web-клиента и worker-сервиса")),
        ],
        [3.4, 6.6, 7.0],
    )

    add_subsection_title(doc, "Структура репозитория")
    add_multi_col_table(
        doc,
        ["Каталог", "Назначение", "Примечание"],
        [
            TableRow(("web/", "Веб-приложение React", "Имеет собственный Dockerfile, nginx.conf и scripts/")),
            TableRow(("mobile/", "Expo / React Native клиент", "Есть android/, APK-артефакт и web-export")),
            TableRow(("docs/techdoc/", "Техническая документация", "Генератор DOCX, ассеты и QA-рендеры")),
        ],
        [3.3, 6.7, 7.0],
    )


def add_architecture(doc: Document) -> None:
    add_section_title(doc, "2. Архитектура решения")
    add_body(
        doc,
        "Архитектурно проект построен вокруг прямого доступа клиентов к Supabase через anon key и вынесенного серверного процесса push-worker, который использует service role key только для безопасной фоновой отправки уведомлений. Такой подход упрощает клиентскую часть и отделяет пользовательские сценарии от фоновых задач.",
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ARCHITECTURE_PATH), width=Inches(6.45))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run("Рисунок 1 — Архитектура проекта Aptechka")
    caption_run.italic = True
    caption_run.font.size = pt(10)
    caption_run.font.color.rgb = MUTED

    add_bullets(
        doc,
        [
            "Веб-клиент и мобильный клиент используют один и тот же Supabase backend и общую предметную модель.",
            "Web часть реализует service worker и хранит push-подписки браузера в таблице push_subscriptions.",
            "Мобильная часть регистрирует Expo push token и сохраняет его в mobile_push_tokens.",
            "push-worker регулярно опрашивает medicine_schedules, отправляет уведомления и пишет факты доставки в push_notification_logs.",
        ],
    )
    add_callout(
        doc,
        "Архитектурное преимущество",
        "Service role key не попадает в клиенты: он используется только в push-worker. Клиенты работают через anon key и стандартные механизмы Supabase Auth, что снижает риск компрометации серверных полномочий.",
    )
    doc.add_page_break()


def add_modules(doc: Document) -> None:
    add_section_title(doc, "3. Функциональные модули и степень готовности")
    add_multi_col_table(
        doc,
        ["Модуль", "Web", "Mobile", "Комментарий"],
        [
            TableRow(("Авторизация и сессии", "Готово", "Готово", "Email/password через Supabase Auth, сохранение сессии на mobile")),
            TableRow(("Создание аптечек", "Готово", "Готово", "Поддерживаются тип, имя, описание и базовые приглашения")),
            TableRow(("Совместный доступ", "Готово", "Готово", "Через aptechka_invites и aptechka_members")),
            TableRow(("Инвентарь лекарств", "Готово", "Готово", "Добавление, удаление, категории, изображения, список по аптечке")),
            TableRow(("Календарь приема", "Готово", "Готово", "Расписания по дням недели, отметка факта приема")),
            TableRow(("Web/browser notifications", "Готово", "Не применимо", "Service worker и browser notification API")),
            TableRow(("Remote push-уведомления", "Готово", "Готово", "Требуют корректных env и SQL-настройки mobile_push_tokens")),
            TableRow(("Профиль пользователя", "Готово", "Частично", "Web позволяет загрузить avatar; mobile profile пока на mock data")),
        ],
        [4.8, 2.2, 2.2, 7.6],
    )

    add_subsection_title(doc, "Основные пользовательские сценарии")
    add_numbered(
        doc,
        [
            "Регистрация пользователя и вход в систему через Supabase Auth.",
            "Создание личной аптечки с выбором типа и описания.",
            "Приглашение другого пользователя по email в совместную аптечку.",
            "Добавление лекарства с дозировкой, категорией, количеством и, при необходимости, изображением из storage.",
            "Настройка расписания приема лекарства по времени и дням недели.",
            "Получение напоминания и фиксация факта приема в medicine_logs.",
        ],
    )
    add_callout(
        doc,
        "Текущее ограничение mobile-клиента",
        "В mobile-приложении навигация пока основана на локальном состоянии экранов, без React Navigation / Expo Router. Для демонстрации это приемлемо, но для промышленной эксплуатации лучше перейти на полноценный навигационный стек.",
        fill=WARN_FILL,
    )
    doc.add_page_break()


def add_data_model(doc: Document) -> None:
    add_section_title(doc, "4. Модель данных и внешние интеграции")
    add_body(
        doc,
        "Полный SQL-дамп всех бизнес-таблиц в репозитории отсутствует, однако состав сущностей надежно восстанавливается по коду клиентов, worker-скрипту и доступному SQL для mobile push. Для защиты этого достаточно, если честно разделять факты из кода и предположения о полном DDL.",
    )

    add_subsection_title(doc, "Бизнес-таблицы Supabase")
    add_multi_col_table(
        doc,
        ["Таблица", "Ключевые поля", "Назначение"],
        [
            TableRow(("aptechkas", "id, owner_id, owner_name, name, description, avatar", "Справочник пользовательских аптечек")),
            TableRow(("aptechka_members", "aptechka_id, user_id, role", "Связь аптечки с приглашенными участниками")),
            TableRow(("aptechka_invites", "id, aptechka_id, invited_email, invited_by, status", "Приглашения в совместные аптечки")),
            TableRow(("medicines", "id, aptechka_id, name, dose, quantity, category, expiry_date, image_url", "Состав аптечек и карточки лекарств")),
            TableRow(("medicine_schedules", "id, user_id, medicine_id, time, days", "Расписания приема по времени и дням недели")),
            TableRow(("medicine_logs", "id, user_id, schedule_id, date, taken", "Журнал фактического приема лекарств")),
            TableRow(("notifications", "id, user_id, type, title, message, is_read", "Лента пользовательских уведомлений и статусы прочтения")),
        ],
        [3.5, 6.3, 6.8],
    )

    add_subsection_title(doc, "Таблицы уведомлений и технические сущности")
    add_multi_col_table(
        doc,
        ["Таблица", "Ключевые поля", "Назначение"],
        [
            TableRow(("push_subscriptions", "user_id, endpoint, p256dh, auth, user_agent", "Браузерные Web Push подписки")),
            TableRow(("push_notification_logs", "user_id, schedule_id, scheduled_for, sent_count", "Защита от повторной отправки напоминаний")),
            TableRow(("mobile_push_tokens", "user_id, expo_push_token, project_id, platform, device_name, updated_at", "Push-токены мобильных устройств; DDL есть в supabase_mobile_push_setup.sql")),
        ],
        [3.8, 6.6, 6.2],
    )

    add_subsection_title(doc, "Bucket-объекты storage")
    add_multi_col_table(
        doc,
        ["Bucket", "Использование", "Комментарий"],
        [
            TableRow(("avatars", "Фотография профиля web-пользователя", "URL сохраняется в user_metadata Supabase Auth")),
            TableRow(("medicine-images", "Изображения лекарств", "Используется web и mobile клиентами при добавлении позиций")),
        ],
        [4.2, 6.6, 5.8],
    )

    add_callout(
        doc,
        "Нюанс по безопасности данных",
        "В репозитории явно показана настройка RLS для mobile_push_tokens. Для остальных таблиц следует отдельно продемонстрировать или описать политики доступа в панели Supabase, если защита будет включать вопросы по безопасности.",
        fill=INFO_FILL,
    )
    doc.add_page_break()


def add_operations(doc: Document) -> None:
    add_section_title(doc, "5. Уведомления, развертывание и эксплуатация")
    add_subsection_title(doc, "Алгоритм напоминаний")
    add_numbered(
        doc,
        [
            "Пользователь создает запись в medicine_schedules и связывает ее с конкретным лекарством.",
            "Веб-клиент обновляет service worker и локальный список расписаний для браузерных напоминаний.",
            "push-worker по таймеру выбирает записи medicine_schedules, совпадающие с текущим временем и днем недели.",
            "Для web используется Web Push по данным из push_subscriptions; для mobile — Expo Push API по токенам из mobile_push_tokens.",
            "После отправки worker записывает факт доставки в push_notification_logs, чтобы одно и то же напоминание не ушло повторно в тот же день.",
            "При подтверждении приема создается запись в medicine_logs и пользовательское уведомление в notifications.",
        ],
    )

    add_subsection_title(doc, "Сервисы Docker Compose")
    add_multi_col_table(
        doc,
        ["Сервис", "Назначение", "Особенности"],
        [
            TableRow(("web", "Статическая публикация React build через nginx", "Использует build stage и монтирует сертификаты Let's Encrypt")),
            TableRow(("push-worker", "Фоновая отправка web/mobile push-напоминаний", "Запускает npm run push:worker и требует service role key")),
        ],
        [3.2, 6.2, 7.2],
    )

    add_subsection_title(doc, "Основные переменные окружения")
    add_multi_col_table(
        doc,
        ["Переменная", "Кому нужна", "Назначение"],
        [
            TableRow(("REACT_APP_SUPABASE_URL", "web", "URL проекта Supabase")),
            TableRow(("REACT_APP_SUPABASE_ANON_KEY", "web", "Публичный клиентский ключ для web SPA")),
            TableRow(("EXPO_PUBLIC_SUPABASE_URL", "mobile", "URL проекта Supabase для Expo")),
            TableRow(("EXPO_PUBLIC_SUPABASE_ANON_KEY", "mobile", "Публичный клиентский ключ для mobile")),
            TableRow(("SUPABASE_SERVICE_ROLE_KEY", "push-worker", "Серверный ключ для фоновой отправки и логирования")),
            TableRow(("VAPID_PUBLIC_KEY / VAPID_PRIVATE_KEY", "web + worker", "Ключи Web Push")),
            TableRow(("PUSH_TIMEZONE", "push-worker", "Таймзона вычисления напоминаний")),
            TableRow(("PUSH_POLL_INTERVAL_MS", "push-worker", "Интервал опроса расписаний")),
        ],
        [5.3, 3.1, 8.2],
    )
    doc.add_page_break()


def add_quality(doc: Document) -> None:
    add_section_title(doc, "6. Подтвержденный статус сборок и технический долг")
    add_body(
        doc,
        f"Локальная проверка от {BUILD_DATE} подтверждает, что проект находится в рабочем состоянии и воспроизводится из репозитория. Проверка выполнялась отдельно для web-клиента и для mobile-клиента в режиме web export.",
    )
    add_multi_col_table(
        doc,
        ["Проверка", "Результат", "Комментарий"],
        [
            TableRow(("aptechka: npm run build", "Успешно", "Production build выполнен, но остались предупреждения ESLint")),
            TableRow(("mobile: npx expo export --platform web", "Успешно", "Сборка подтверждена после запуска вне sandbox")),
            TableRow(("mobile APK артефакт", "Доступен", "В репозитории присутствует mobile/dist-apk/Aptechka-Mobile-v1.0.0-debug.apk")),
        ],
        [5.6, 2.8, 8.2],
    )

    add_subsection_title(doc, "Что стоит проговорить как технический долг")
    add_bullets(
        doc,
        [
            "В web-клиенте есть предупреждения ESLint: неиспользуемые импорты и missing dependencies в useEffect.",
            "Профиль mobile-клиента пока оформлен как демонстрационный экран с mock data.",
            "Полный SQL-слой всех бизнес-таблиц не зафиксирован в репозитории отдельной миграцией.",
            "Лендинг говорит о записи к врачу, но кода реального сценария записи в текущем проекте нет.",
            "Навигация mobile-клиента реализована локальным состоянием и требует дальнейшего развития до production-grade маршрутизации.",
        ],
    )

    add_callout(
        doc,
        "Сильная сторона проекта",
        "Несмотря на перечисленный технический долг, базовая бизнес-вертикаль уже связана end-to-end: авторизация, совместные аптечки, лекарства, расписание, журнал приема и доставка напоминаний. Для защиты это весомый аргумент, потому что демонстрируется полный пользовательский цикл.",
        fill=OK_FILL,
    )


def add_defense_script(doc: Document) -> None:
    add_section_title(doc, "7. Рекомендуемый сценарий демонстрации на защите")
    add_numbered(
        doc,
        [
            "Показать лендинг и кратко сформулировать проблему: пользователю сложно контролировать состав аптечки и график приема лекарств.",
            "Войти в систему и открыть раздел инвентаря.",
            "Создать новую аптечку, выбрать тип и при необходимости пригласить второго пользователя по email.",
            "Добавить лекарство с параметрами дозировки и категории; при наличии — загрузить изображение.",
            "Перейти в календарь, создать расписание приема по времени и дням недели.",
            "Показать фиксацию факта приема и появление записи в пользовательских уведомлениях.",
            "Если окружение подготовлено, продемонстрировать push-уведомление в web или mobile.",
            "Завершить архитектурным слайдом: один Supabase backend, два клиента и отдельный worker уведомлений.",
        ],
    )

    add_callout(
        doc,
        "Как лучше позиционировать проект",
        "Проект стоит защищать как прикладовую цифровую систему учета лекарств с функциями совместного доступа и автоматических напоминаний. Это честная и технологически подтвержденная формулировка текущего результата.",
        fill=LIGHT_BG,
    )


def build_document() -> Path:
    ensure_dirs()
    create_architecture_diagram()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_overview(doc)
    add_architecture(doc)
    add_modules(doc)
    add_data_model(doc)
    add_operations(doc)
    add_quality(doc)
    add_defense_script(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output_path = build_document()
    print(output_path)
